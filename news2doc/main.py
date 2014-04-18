#!/usr/bin/env python
# encoding: utf-8

from wsj import *
from ft import *
from nyt import *

# Use pip install python-docx to install dependency.
# https://python-docx.readthedocs.org/en/latest/
from docx import *

def make_doc(articles):
	'''
	Make Microsoft Word document from fetched articles.
	'''
	print 'Making .docx file...'

	document = Document()
	for article in articles:
		document.add_heading(article['title'], 1)
		document.add_paragraph(article['link'])
		document.add_paragraph(str(article['publish_date']))
		document.add_paragraph()
		if 'text' not in article: return
		text = article['text']
		if not text: return
		for line in text:
			try:
				line = line.decode('utf-8', 'ignore')
			except:
				pass
			line = line.strip()
			if len(line):
				document.add_paragraph(line)
		document.add_paragraph()

	document.save('news.docx')
	print('...Done.')


def main():
	articles = []
	try: articles += WSJ().fetch_all_articles_in_24hours()
	except: pass
	try: articles += FinancialTimes().fetch_all_articles_in_24hours()
	except: pass
	try: articles += NewYorkTimes().fetch_all_articles_in_24hours()
	except: pass
	make_doc(articles)

if __name__ == '__main__':
	main()
